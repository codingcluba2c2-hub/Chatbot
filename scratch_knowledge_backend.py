import os

os.makedirs('backend/services/knowledge', exist_ok=True)
os.makedirs('backend/uploads', exist_ok=True)

with open('backend/schemas/knowledge.py', 'w') as f:
    f.write('''\
from pydantic import BaseModel, Field
from typing import Optional, List, Dict, Any
import uuid
import time

class KnowledgeDocument(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    category: str = "general"
    language: str = "en"
    status: str = "processing" # draft, processing, published, archived, failed
    version: int = 1
    tags: List[str] = Field(default_factory=list)
    uploaded_by: str = "admin"
    created_at: float = Field(default_factory=time.time)
    updated_at: float = Field(default_factory=time.time)
    file_path: Optional[str] = None
    file_type: Optional[str] = None
    file_size: Optional[int] = 0
    chunk_count: int = 0
    error_message: Optional[str] = None

class DocumentChunk(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    content: str
    metadata: Dict[str, Any] = Field(default_factory=dict)
    created_at: float = Field(default_factory=time.time)
''')

with open('backend/services/knowledge/chunking_engine.py', 'w') as f:
    f.write('''\
import re
from typing import List

class ChunkingEngine:
    @staticmethod
    def chunk_text(text: str, strategy: str = "character", max_chars: int = 1000, overlap: int = 200) -> List[str]:
        if not text:
            return []
            
        chunks = []
        if strategy == "character":
            start = 0
            text_len = len(text)
            while start < text_len:
                end = min(start + max_chars, text_len)
                
                # If we're not at the end, try to break at a newline or space
                if end < text_len:
                    last_newline = text.rfind('\\n', start, end)
                    last_space = text.rfind(' ', start, end)
                    
                    if last_newline != -1 and last_newline > start + overlap:
                        end = last_newline
                    elif last_space != -1 and last_space > start + overlap:
                        end = last_space
                        
                chunks.append(text[start:end].strip())
                start = end - overlap if end < text_len else end
                
        elif strategy == "paragraph":
            # Simple split by double newline
            paragraphs = re.split(r'\\n\\s*\\n', text)
            current_chunk = ""
            for p in paragraphs:
                p = p.strip()
                if not p:
                    continue
                if len(current_chunk) + len(p) < max_chars:
                    current_chunk += "\\n\\n" + p if current_chunk else p
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = p
            if current_chunk:
                chunks.append(current_chunk)
                
        return chunks
''')

with open('backend/services/knowledge/extraction_engine.py', 'w') as f:
    f.write('''\
import os
import io
from core.logger import get_logger

logger = get_logger(__name__)

class ExtractionEngine:
    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        ext = file_type.lower()
        if ext == 'txt' or ext == 'md' or ext == 'csv':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif ext == 'pdf':
            try:
                import PyPDF2
                text = ""
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\\n"
                return text
            except ImportError:
                logger.warning("PyPDF2 not installed. Returning empty text for PDF.")
                return "[PDF Extraction Requires PyPDF2]"
        elif ext == 'docx':
            try:
                import docx
                doc = docx.Document(file_path)
                return "\\n".join([p.text for p in doc.paragraphs])
            except ImportError:
                logger.warning("python-docx not installed. Returning empty text for DOCX.")
                return "[DOCX Extraction Requires python-docx]"
        
        # Fallback
        return "[Unsupported File Type]"
''')

with open('backend/api/routes/knowledge.py', 'w') as f:
    f.write('''\
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks
from typing import List, Optional
import shutil
import os
import time
from schemas.knowledge import KnowledgeDocument, DocumentChunk
from repositories.registry import document_repo, chunk_repo, log_audit
from services.knowledge.extraction_engine import ExtractionEngine
from services.knowledge.chunking_engine import ChunkingEngine

router = APIRouter(prefix="/api/knowledge", tags=["knowledge"])

UPLOAD_DIR = "uploads"

def process_document_bg(doc_id: str, file_path: str, file_type: str):
    doc = document_repo.get(doc_id)
    if not doc:
        return
        
    try:
        # Extract text
        text = ExtractionEngine.extract_text(file_path, file_type)
        
        # Chunk text
        chunks = ChunkingEngine.chunk_text(text, strategy="paragraph")
        
        # Save chunks
        for i, chunk_text in enumerate(chunks):
            chunk = DocumentChunk(
                document_id=doc_id,
                content=chunk_text,
                metadata={"index": i, "token_estimate": len(chunk_text) // 4}
            )
            chunk_repo.create(chunk)
            
        # Update doc status
        doc.status = "published"
        doc.chunk_count = len(chunks)
        doc.updated_at = time.time()
        document_repo.update(doc_id, doc)
        
        log_audit("process", "KnowledgeDocument", doc_id, new_value={"status": "published", "chunks": len(chunks)})
        
    except Exception as e:
        doc.status = "failed"
        doc.error_message = str(e)
        doc.updated_at = time.time()
        document_repo.update(doc_id, doc)
        log_audit("process_failed", "KnowledgeDocument", doc_id, new_value={"error": str(e)})

@router.post("/upload", response_model=KnowledgeDocument)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    category: str = Form("general"),
    tags: str = Form(""),
    language: str = Form("en")
):
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    
    file_ext = file.filename.split('.')[-1].lower() if '.' in file.filename else 'txt'
    safe_filename = f"{time.time()}_{file.filename}"
    file_path = os.path.join(UPLOAD_DIR, safe_filename)
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    file_size = os.path.getsize(file_path)
    tag_list = [t.strip() for t in tags.split(",") if t.strip()]
    
    doc = KnowledgeDocument(
        name=file.filename,
        category=category,
        language=language,
        tags=tag_list,
        file_path=file_path,
        file_type=file_ext,
        file_size=file_size,
        status="processing"
    )
    
    created = document_repo.create(doc)
    log_audit("create", "KnowledgeDocument", created.id, new_value=created.model_dump())
    
    # Process in background
    background_tasks.add_task(process_document_bg, created.id, file_path, file_ext)
    
    return created

@router.get("/documents", response_model=List[KnowledgeDocument])
def list_documents(skip: int = 0, limit: int = 100):
    return document_repo.get_all(skip=skip, limit=limit)

@router.delete("/documents/{doc_id}")
def delete_document(doc_id: str):
    doc = document_repo.get(doc_id)
    if doc:
        # Delete chunks
        all_chunks = chunk_repo.get_all()
        chunks_to_delete = [c for c in all_chunks if c.document_id == doc_id]
        for c in chunks_to_delete:
            chunk_repo.delete(c.id)
            
        # Delete file
        if doc.file_path and os.path.exists(doc.file_path):
            os.remove(doc.file_path)
            
        document_repo.delete(doc_id)
        log_audit("delete", "KnowledgeDocument", doc_id)
        return {"ok": True}
    raise HTTPException(404, "Not found")

@router.get("/documents/{doc_id}/chunks", response_model=List[DocumentChunk])
def get_document_chunks(doc_id: str):
    all_chunks = chunk_repo.get_all()
    return [c for c in all_chunks if c.document_id == doc_id]
''')
